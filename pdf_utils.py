import io

import pymupdf
from docx import Document
from docx.shared import Pt


def extract_text_from_pdf(file_bytes: bytes) -> list[dict]:
    """Extract paragraph-like text blocks from a text-based PDF, in reading order.

    Returns list of dicts compatible with the pptx_utils text_unit schema,
    with 'slide_idx' holding the PDF page index.
    """
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    units = []
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        blocks = [b for b in page.get_text("dict")["blocks"] if b.get("type") == 0]
        blocks.sort(key=lambda b: (round(b["bbox"][1], 1), b["bbox"][0]))
        for b_idx, block in enumerate(blocks):
            lines, sizes = [], []
            for line in block.get("lines", []):
                line_text = "".join(span["text"] for span in line.get("spans", []))
                if line_text.strip():
                    lines.append(line_text)
                sizes.extend(span.get("size", 12.0) for span in line.get("spans", []))
            text = "\n".join(lines).strip()
            if not text:
                continue
            units.append({
                "id": f"pdf_p{page_idx}_b{b_idx}",
                "ko_text": text,
                "slide_idx": page_idx,
                "shape_id": b_idx,
                "p_idx": 0,
                "font_size": round(sum(sizes) / len(sizes), 1) if sizes else 12.0,
                "shape_text": text,
                "shape_para_count": len(lines) or 1,
            })
    doc.close()
    return units


def page_count(file_bytes: bytes) -> int:
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    n = doc.page_count
    doc.close()
    return n


def render_pdf_to_images(file_bytes: bytes, dpi: int = 100) -> list[bytes]:
    """Render each PDF page to a PNG (for slide-style previews). Returns [] on failure."""
    try:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        zoom = dpi / 72.0
        matrix = pymupdf.Matrix(zoom, zoom)
        images = [page.get_pixmap(matrix=matrix).tobytes("png") for page in doc]
        doc.close()
        return images
    except Exception:
        return []


def build_translated_docx(units: list[dict], translations: dict[str, str]) -> bytes:
    """Build a new Word document from translated PDF text units, grouped by page."""
    doc = Document()
    by_page: dict[int, list[dict]] = {}
    for u in units:
        by_page.setdefault(u["slide_idx"], []).append(u)

    for page_idx in sorted(by_page.keys()):
        if page_idx > 0:
            doc.add_page_break()
        for u in by_page[page_idx]:
            text = translations.get(u["id"], "").strip() or u["ko_text"]
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(14) if u.get("font_size", 12.0) >= 14 else Pt(11)
            if u.get("font_size", 12.0) >= 14:
                run.bold = True

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
