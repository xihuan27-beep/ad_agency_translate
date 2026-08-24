import io
import time
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> list[dict]:
    """Extract text units from a Word (.docx) document.

    Returns list of dicts compatible with the pptx_utils text_unit schema,
    using 'ko_text' as the source text field (direction-agnostic naming).
    """
    doc = Document(io.BytesIO(file_bytes))
    units = []
    ts = int(time.time() * 1000)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        units.append({
            "id": f"doc_{ts}_p{para_idx}",
            "ko_text": text,
            "slide_idx": 0,
            "shape_id": para_idx,
            "p_idx": 0,
            "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 12.0,
            "shape_text": text,
            "shape_para_count": 1,
        })

    for tbl_idx, table in enumerate(doc.tables):
        seen_cells = set()
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                text = cell.text.strip()
                if not text:
                    continue
                units.append({
                    "id": f"doc_{ts}_t{tbl_idx}r{row_idx}c{cell_idx}",
                    "ko_text": text,
                    "slide_idx": 0,
                    "shape_id": tbl_idx * 10000 + row_idx * 100 + cell_idx,
                    "p_idx": 0,
                    "font_size": 12.0,
                    "shape_text": text,
                    "shape_para_count": 1,
                })

    return units


def apply_translations_to_docx(file_bytes: bytes, translations: dict[str, str]) -> bytes:
    """Replace source text in a .docx with translated text, preserving formatting."""
    doc = Document(io.BytesIO(file_bytes))

    def _replace_para(para, new_text: str):
        if not para.runs:
            para.add_run(new_text)
            return
        # Copy formatting from first run; clear others
        first = para.runs[0]
        first.text = new_text
        for run in para.runs[1:]:
            run.text = ""

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # Match by position — IDs are generated deterministically so we search
        # translations by looking for any key that ends with f"_p{para_idx}"
        matched = next((v for k, v in translations.items() if k.endswith(f"_p{para_idx}")), None)
        if matched and matched.strip():
            _replace_para(para, matched)

    for tbl_idx, table in enumerate(doc.tables):
        seen_cells = set()
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                text = cell.text.strip()
                if not text:
                    continue
                key_suffix = f"_t{tbl_idx}r{row_idx}c{cell_idx}"
                matched = next((v for k, v in translations.items() if k.endswith(key_suffix)), None)
                if matched and matched.strip():
                    for para in cell.paragraphs:
                        _replace_para(para, "")
                    _replace_para(cell.paragraphs[0], matched)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
